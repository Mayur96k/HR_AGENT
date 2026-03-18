"""
file_reader.py — Read PDF and DOCX files to plain text
Steps 1 & 2: Read files from folder, convert to text
"""
import logging
from pathlib import Path

log = logging.getLogger(__name__)


def read_file_to_text(file_path) -> str:
    """Read PDF or DOCX file and return plain text."""
    path = Path(file_path)
    suffix = path.suffix.lower()

    try:
        if suffix == ".pdf":
            return _read_pdf(path)
        elif suffix in (".docx", ".doc"):
            return _read_docx(path)
        else:
            log.warning(f"Unsupported format: {suffix}")
            return ""
    except Exception as e:
        log.error(f"Failed to read {path.name}: {e}")
        return ""


def _read_pdf(path: Path) -> str:
    import pdfplumber
    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)
    return "\n".join(text_parts).strip()


def _read_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also grab table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs).strip()
